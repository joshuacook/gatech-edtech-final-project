# api/jobs/assets/tables.py
import csv
import json
import logging
import os
from typing import Dict, List, Tuple

from docx import Document
from jobs.assets.base import AssetProcessor
from langfuse.decorators import observe
from utils.db_utils import update_asset_status

logger = logging.getLogger(__name__)


class TableProcessor(AssetProcessor):
    """Process tables from DOCX files"""

    processor_type = "tables"
    dependencies = []  # No dependencies - works directly with raw file

    def __init__(self, file_hash: str):
        super().__init__(file_hash)

    @observe(name="asset_processor_tables")
    def process(self):
        """Main processing method"""
        try:
            # Check if file is a DOCX
            if (
                self.asset["file_type"]
                != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                logger.info(
                    f"Asset {self.file_hash} is not a DOCX file, skipping table processing"
                )
                update_asset_status(self.file_hash, "tables_skipped")
                return True

            update_asset_status(self.file_hash, "processing_tables")
            logger.info(f"Starting table processing for {self.file_hash}")

            # Create the Document object
            doc = Document(self.get_raw_file_path())

            # Create tables directory
            tables_dir = os.path.join(self.processed_dir, "tables")
            os.makedirs(tables_dir, exist_ok=True)

            # Process the tables
            table_paths, tables_meta = self._process_tables(doc, tables_dir)

            # Update asset record
            update_data = {
                "has_tables": len(table_paths) > 0,
                "table_count": len(table_paths),
                "processed_paths.tables": table_paths,
            }
            self._update_asset(update_data)

            logger.info(
                f"Completed table processing for {self.file_hash}: found {len(table_paths)} tables"
            )
            update_asset_status(self.file_hash, "tables_complete")

            return True

        except Exception as e:
            logger.error(f"Error processing tables for {self.file_hash}: {str(e)}")
            update_asset_status(self.file_hash, "tables_error", str(e))
            raise

    def _extract_table(self, table) -> List[List[str]]:
        """Convert a table to a list of lists with improved text extraction"""
        try:
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    # Clean and normalize cell text
                    text = cell.text.strip()
                    # Replace multiple spaces/newlines with single space
                    text = " ".join(text.split())
                    row_data.append(text)

                # Only add non-empty rows
                if any(cell for cell in row_data):
                    table_data.append(row_data)

            return table_data

        except Exception as e:
            logger.error(f"Error extracting table data: {str(e)}")
            raise

    def _process_tables(
        self, doc: Document, tables_dir: str
    ) -> Tuple[Dict[str, dict], Dict[str, dict]]:
        """Process all tables from document"""
        table_paths = {}
        tables_meta = {}

        try:
            for i, table in enumerate(doc.tables):
                try:
                    table_name = f"table_{i}"
                    logger.debug(f"Processing {table_name}")

                    # Extract and validate table data
                    table_data = self._extract_table(table)

                    if not table_data or not any(
                        any(cell for cell in row) for row in table_data
                    ):
                        logger.debug(f"Skipping empty table: {table_name}")
                        continue

                    # Save table as CSV
                    csv_path = os.path.join(tables_dir, f"{table_name}.csv")
                    with open(csv_path, "w", newline="", encoding="utf-8") as f:
                        writer = csv.writer(f)
                        writer.writerows(table_data)

                    # Generate and save table HTML
                    html_path = os.path.join(tables_dir, f"{table_name}.html")
                    self._save_table_html(table_data, html_path)

                    # Store paths and metadata
                    table_paths[table_name] = {"csv": csv_path, "html": html_path}

                    tables_meta[table_name] = {
                        "num_rows": len(table_data),
                        "num_cols": len(table_data[0]) if table_data else 0,
                        "headers": table_data[0] if table_data else [],
                        "empty_cells": self._count_empty_cells(table_data),
                        "total_cells": sum(len(row) for row in table_data),
                    }

                except Exception as e:
                    logger.error(f"Error processing table {i}: {str(e)}")
                    continue

            # Save metadata
            meta_path = os.path.join(tables_dir, "tables_meta.json")
            with open(meta_path, "w", encoding="utf-8") as f:
                json.dump(tables_meta, f, ensure_ascii=False, indent=2)
                logger.debug(f"Saved table metadata to {meta_path}")

            return table_paths, tables_meta

        except Exception as e:
            logger.error(f"Error in _process_tables: {str(e)}")
            raise

    def _save_table_html(self, table_data: List[List[str]], html_path: str):
        """Generate and save HTML representation of table"""
        try:
            html = ['<table border="1" class="table">']

            # Add header row
            if table_data:
                html.append("<thead><tr>")
                for header in table_data[0]:
                    html.append(f"<th>{header}</th>")
                html.append("</tr></thead>")

            # Add data rows
            html.append("<tbody>")
            for row in table_data[1:]:
                html.append("<tr>")
                for cell in row:
                    html.append(f"<td>{cell}</td>")
                html.append("</tr>")
            html.append("</tbody>")
            html.append("</table>")

            with open(html_path, "w", encoding="utf-8") as f:
                f.write("\n".join(html))

        except Exception as e:
            logger.error(f"Error saving table HTML: {str(e)}")
            raise

    def _count_empty_cells(self, table_data: List[List[str]]) -> int:
        """Count number of empty cells in table"""
        return sum(1 for row in table_data for cell in row if not cell.strip())
