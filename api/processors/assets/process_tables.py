import json
import logging
import os
from typing import Tuple

from docx import Document
from processors.base import BaseAssetProcessor

logger = logging.getLogger(__name__)


class ProcessTables(BaseAssetProcessor):
    def __init__(self):
        super().__init__("tables", "table", requires_docx=True)

    async def process_asset(self, file_hash: str, asset: dict, db: dict, span):
        processed_dir = self.get_processed_dir(file_hash)
        tables_dir = os.path.join(processed_dir, "tables")
        os.makedirs(tables_dir, exist_ok=True)

        doc = Document(os.path.join("/app/filestore/raw", f"{file_hash}.docx"))
        tables_paths = {}
        tables_meta = {}

        for i, table in enumerate(doc.tables):
            try:
                table_name, table_paths, table_meta = self._process_table(
                    i, table, tables_dir
                )
                tables_paths[table_name] = table_paths
                tables_meta[table_name] = table_meta

            except Exception as e:
                logger.error(f"Error processing table {i}: {str(e)}")
                continue

        meta_path = os.path.join(tables_dir, "tables_meta.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(tables_meta, f, ensure_ascii=False, indent=2)

        update_data = {
            "has_tables": len(tables_paths) > 0,
            "table_count": len(tables_paths),
            "processed_paths.tables": tables_paths,
        }
        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})
        return {
            "status": "success",
            "table_count": len(tables_paths),
            "table_paths": tables_paths,
            "tables_meta": tables_meta,
        }

    def _process_table(self, i: int, table, tables_dir: str) -> Tuple[str, dict]:
        table_name = f"table_{i}"
        table_data = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()
                text = " ".join(text.split())
                row_data.append(text)
            if any(cell for cell in row_data):
                table_data.append(row_data)

        csv_path = self._write_table_to_csv(table_name, table_data, tables_dir)
        html_path = self._write_table_to_html(table_name, table_data, tables_dir)

        table_paths = {"csv": csv_path, "html": html_path}
        table_meta = {
            "num_rows": len(table_data),
            "num_cols": len(table_data[0]) if table_data else 0,
            "headers": table_data[0] if table_data else [],
            "empty_cells": sum(
                1 for row in table_data for cell in row if not cell.strip()
            ),
            "total_cells": sum(len(row) for row in table_data),
        }

        if not table_data:
            return None, None, None
        return table_name, table_paths, table_meta

    def _write_table_to_csv(self, table_name: str, table_data: dict, tables_dir: str):
        csv_path = os.path.join(tables_dir, f"{table_name}.csv")
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            import csv

            writer = csv.writer(f)
            writer.writerows(table_data)
        return csv_path

    def _write_table_to_html(self, table_name: str, table_data: dict, tables_dir: str):
        html_path = os.path.join(tables_dir, f"{table_name}.html")
        with open(html_path, "w", encoding="utf-8") as f:
            html = ['<table border="1" class="table">']
            if table_data:
                html.append("<thead><tr>")
                for header in table_data[0]:
                    html.append(f"<th>{header}</th>")
                html.append("</tr></thead>")
            html.append("<tbody>")
            for row in table_data[1:]:
                html.append("<tr>")
                for cell in row:
                    html.append(f"<td>{cell}</td>")
                html.append("</tr>")
            html.append("</tbody></table>")
            f.write("\n".join(html))
        return html_path
