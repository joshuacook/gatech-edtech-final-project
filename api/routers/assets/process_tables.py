# api/routers/assets.py

import json
import logging
import os
from typing import Dict, Optional

from docx import Document
from fastapi import APIRouter, Header, HTTPException
from langfuse import Langfuse
from utils.db_utils import init_mongo, update_asset_status
from utils.langfuse_utils import configure_langfuse

logger = logging.getLogger(__name__)
process_tables_router = APIRouter()

configure_langfuse()
langfuse = Langfuse()


@process_tables_router.post("/process_tables/{file_hash}")
async def process_tables(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Process tables from a DOCX file"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        if (
            asset["file_type"]
            != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            logger.info(
                f"Asset {file_hash} is not a DOCX file, skipping table processing"
            )
            update_asset_status(file_hash, "tables_skipped")
            return {"status": "skipped", "reason": "not_docx"}

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="table_processing",
            metadata={"processor_type": "tables", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_tables")
        logger.info(f"Starting table processing for {file_hash}")

        processed_dir = os.path.join("/app/filestore/processed", file_hash)
        tables_dir = os.path.join(processed_dir, "tables")
        os.makedirs(tables_dir, exist_ok=True)

        doc = Document(os.path.join("/app/filestore/raw", f"{file_hash}.docx"))
        table_paths = {}
        tables_meta = {}

        for i, table in enumerate(doc.tables):
            try:
                table_name = f"table_{i}"
                table_data = []

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        text = " ".join(text.split())
                        row_data.append(text)
                    if any(cell for cell in row_data):
                        table_data.append(row_data)

                if not table_data:
                    continue

                # Save as CSV
                csv_path = os.path.join(tables_dir, f"{table_name}.csv")
                with open(csv_path, "w", newline="", encoding="utf-8") as f:
                    import csv

                    writer = csv.writer(f)
                    writer.writerows(table_data)

                # Save as HTML
                html_path = os.path.join(tables_dir, f"{table_name}.html")
                with open(html_path, "w", encoding="utf-8") as f:
                    html = ['<table border="1" class="table">']
                    if table_data:
                        html.append("<thead><tr>")
                        for header in table_data[0]:
                            html.append(f"<th>{header}</th>")
                        html.append("</tr></thead>")
                    html.append("<tbody>")
                    for row in table_data[1:]:
                        html.append("<tr>")
                        for cell in row:
                            html.append(f"<td>{cell}</td>")
                        html.append("</tr>")
                    html.append("</tbody></table>")
                    f.write("\n".join(html))

                table_paths[table_name] = {"csv": csv_path, "html": html_path}

                # Collect metadata
                tables_meta[table_name] = {
                    "num_rows": len(table_data),
                    "num_cols": len(table_data[0]) if table_data else 0,
                    "headers": table_data[0] if table_data else [],
                    "empty_cells": sum(
                        1 for row in table_data for cell in row if not cell.strip()
                    ),
                    "total_cells": sum(len(row) for row in table_data),
                }

            except Exception as e:
                logger.error(f"Error processing table {i}: {str(e)}")
                continue

        # Save metadata
        meta_path = os.path.join(tables_dir, "tables_meta.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(tables_meta, f, ensure_ascii=False, indent=2)

        # Update asset record
        update_data = {
            "has_tables": len(table_paths) > 0,
            "table_count": len(table_paths),
            "processed_paths.tables": table_paths,
        }
        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        logger.info(f"Completed table processing for {file_hash}")
        update_asset_status(file_hash, "tables_complete")

        span.event(
            name="table_processing_complete",
            metadata={"table_count": len(table_paths)},
        )

        return {
            "status": "success",
            "table_count": len(table_paths),
            "table_paths": table_paths,
            "tables_meta": tables_meta,
        }

    except Exception as e:
        error_msg = f"Error processing tables: {str(e)}"
        if "span" in locals():
            span.event(
                name="table_processing_error",
                metadata={"error": error_msg},
                level="error",
            )
        update_asset_status(file_hash, "tables_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()
