# api/routers/assets.py

import hashlib
import json
import logging
import os
import time
import zipfile
from datetime import datetime
from typing import Dict, Optional

import requests
from docx import Document
from fastapi import APIRouter, Header, HTTPException
from langfuse import Langfuse
from utils import handle_error
from utils.db_utils import init_mongo, update_asset_status
from utils.image_utils import format_image_info, get_image_technical_metadata
from utils.langfuse_utils import configure_langfuse
from utils.lexeme_utils import get_prompts_for_category, merge_lexeme_results

from .chat import chat_call

logger = logging.getLogger(__name__)
assets_router = APIRouter(prefix="/assets")

configure_langfuse()
langfuse = Langfuse()


@assets_router.post("/process_images/{file_hash}")
async def process_images(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Process images from a DOCX file"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        if (
            asset["file_type"]
            != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            logger.info(
                f"Asset {file_hash} is not a DOCX file, skipping image processing"
            )
            update_asset_status(file_hash, "images_skipped")
            return {"status": "skipped", "reason": "not_docx"}

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="image_processing",
            metadata={"processor_type": "images", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_images")
        logger.info(f"Starting image processing for {file_hash}")

        # Create images directory
        processed_dir = os.path.join("/app/filestore/processed", file_hash)
        images_dir = os.path.join(processed_dir, "images")
        os.makedirs(images_dir, exist_ok=True)

        # Process images
        image_paths = {}
        valid_extensions = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}

        with zipfile.ZipFile(
            os.path.join("/app/filestore/raw", f"{file_hash}.docx")
        ) as doc_zip:
            media_files = [f for f in doc_zip.namelist() if f.startswith("word/media/")]

            for file_path in media_files:
                _, ext = os.path.splitext(file_path.lower())
                if ext not in valid_extensions:
                    continue

                image_data = doc_zip.read(file_path)
                image_hash = hashlib.md5(image_data).hexdigest()
                unique_filename = f"{image_hash}{ext}"
                image_path = os.path.join(images_dir, unique_filename)

                if not os.path.exists(image_path):
                    with open(image_path, "wb") as f:
                        f.write(image_data)

                original_name = os.path.basename(file_path)
                image_paths[original_name] = image_path

        # Update asset record
        update_data = {
            "has_images": len(image_paths) > 0,
            "image_count": len(image_paths),
            "processed_paths.images": image_paths,
        }
        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        logger.info(f"Completed image processing for {file_hash}")
        update_asset_status(file_hash, "images_complete")

        span.event(
            name="image_processing_complete",
            metadata={"image_count": len(image_paths)},
        )

        return {
            "status": "success",
            "image_count": len(image_paths),
            "image_paths": image_paths,
        }

    except Exception as e:
        error_msg = f"Error processing images: {str(e)}"
        if "span" in locals():
            span.event(
                name="image_processing_error",
                metadata={"error": error_msg},
                level="error",
            )
        update_asset_status(file_hash, "images_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()


@assets_router.post("/process_image_metadata/{file_hash}")
async def process_image_metadata(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Process metadata for extracted images"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        processed_paths = asset.get("processed_paths", {})
        images = processed_paths.get("images", {})

        if not images:
            logger.info(f"No images found for {file_hash}, skipping metadata")
            update_asset_status(file_hash, "image_metadata_skipped")
            return {"status": "skipped", "reason": "no_images"}

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="image_metadata_processing",
            metadata={"processor_type": "image_metadata", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_image_metadata")
        logger.info(f"Starting image metadata processing for {file_hash}")

        # Read prompt template
        prompts_dir = os.path.join("/app", "prompts", "assets", "image")
        with open(os.path.join(prompts_dir, "metadata.txt"), "r") as f:
            prompt_template = f.read()

        image_metadata = {}

        for image_name, image_path in images.items():
            try:
                # Get technical metadata
                technical_metadata = get_image_technical_metadata(
                    image_name, image_path
                )
                if not technical_metadata:
                    continue

                # Convert image info to text description
                image_description = format_image_info(technical_metadata)
                prompt = f"{prompt_template}\n\nImage Information:\n{image_description}"

                response = chat_call(query=prompt, messages=[])

                if "error" in response:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Chat API error for image {image_name} with error: {response['error']}",
                    )

                response_text = response["message"]
                json_start = response_text.find("{")
                json_end = response_text.rfind("}") + 1

                if json_start < 0 or json_end <= json_start:
                    logger.error(
                        f"Could not find JSON in chat response for image {image_name}"
                    )
                    continue

                ai_metadata = json.loads(response_text[json_start:json_end])
                image_metadata[image_name] = {
                    "technical": technical_metadata,
                    "analysis": ai_metadata,
                }

            except Exception as e:
                logger.error(f"Error processing image {image_name}: {str(e)}")
                continue

        if image_metadata:
            metadata_path = os.path.join(
                "/app/filestore/processed", file_hash, "images", "image_metadata.json"
            )
            os.makedirs(os.path.dirname(metadata_path), exist_ok=True)

            with open(metadata_path, "w", encoding="utf-8") as f:
                json.dump(image_metadata, f, ensure_ascii=False, indent=2)

            update_data = {
                "processed_paths.image_metadata": metadata_path,
                "image_metadata": image_metadata,
            }
            db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        update_asset_status(file_hash, "image_metadata_complete")
        span.event(
            name="image_metadata_complete",
            metadata={"image_count": len(image_metadata)},
        )

        return {
            "status": "success",
            "image_metadata": image_metadata,
            "image_count": len(image_metadata),
        }

    except Exception as e:
        error_msg = f"Error processing image metadata: {str(e)}"
        if "span" in locals():
            span.event(
                name="image_metadata_error",
                metadata={"error": error_msg},
                level="error",
            )
        update_asset_status(file_hash, "image_metadata_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()


@assets_router.post("/process_lexemes/{file_hash}")
async def process_lexemes(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Extract lexemes from document content"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="lexeme_processing",
            metadata={"processor_type": "lexemes", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_lexemes")
        logger.info(f"Starting lexeme processing for {file_hash}")

        # Get the content and metadata
        processed_paths = asset.get("processed_paths", {})
        if not processed_paths or "markdown" not in processed_paths:
            raise HTTPException(
                status_code=400, detail="Markdown content not available"
            )

        with open(processed_paths["markdown"], "r") as f:
            content = f.read()

        category = (
            asset.get("metadata", {})
            .get("documentMetadata", {})
            .get("primaryType", {})
            .get("category", "General/Mixed")
        )
        prompts_to_run = get_prompts_for_category(category)
        all_lexemes = []

        for prompt_file in prompts_to_run:
            try:
                prompt_path = os.path.join(
                    "/app", "prompts", "assets", "lexeme", prompt_file
                )
                with open(prompt_path, "r") as f:
                    prompt_template = f.read()

                prompt = f"{prompt_template}\n\nDocument Content:\n{content}"

                generation = span.generation(
                    name=f"lexeme_extraction_{prompt_file}",
                    input=prompt,
                )

                response = chat_call(query=prompt, messages=[])

                if "error" in response:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Chat API error for prompt {prompt_file} with error: {response['error']}",
                    )

                response_text = response["message"]
                json_start = response_text.find("{")
                json_end = response_text.rfind("}") + 1

                if json_start < 0 or json_end <= json_start:
                    raise ValueError(
                        f"Could not find JSON in chat response for prompt {prompt_file}"
                    )

                prompt_lexemes = json.loads(response_text[json_start:json_end])[
                    "lexemes"
                ]
                all_lexemes.extend(prompt_lexemes)

                generation.end(
                    output=response_text,
                    metadata={
                        "prompt_file": prompt_file,
                        "lexeme_count": len(prompt_lexemes),
                    },
                )

            except Exception as e:
                logger.error(f"Error processing with prompt {prompt_file}: {str(e)}")
                continue

        # Merge and deduplicate lexemes
        merged_lexemes = merge_lexeme_results(all_lexemes)

        # Save results
        lexemes_path = os.path.join(
            "/app/filestore/processed", file_hash, "lexemes.json"
        )
        os.makedirs(os.path.dirname(lexemes_path), exist_ok=True)

        lexeme_data = {
            "lexemes": merged_lexemes,
            "metadata": {
                "count": len(merged_lexemes),
                "processed_at": datetime.now().isoformat(),
                "asset_id": file_hash,
            },
        }

        with open(lexemes_path, "w", encoding="utf-8") as f:
            json.dump(lexeme_data, f, ensure_ascii=False, indent=2)

        # Update asset
        update_data = {
            "processed_paths.lexemes": lexemes_path,
            "lexemes": merged_lexemes,
            "lexeme_count": len(merged_lexemes),
        }
        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        update_asset_status(file_hash, "lexemes_complete")
        span.event(
            name="lexemes_complete", metadata={"lexeme_count": len(merged_lexemes)}
        )

        return {
            "status": "success",
            "lexeme_count": len(merged_lexemes),
            "lexemes": merged_lexemes,
        }

    except Exception as e:
        error_msg = f"Error processing lexemes: {str(e)}"
        if "span" in locals():
            span.event(
                name="lexemes_error", metadata={"error": error_msg}, level="error"
            )
        update_asset_status(file_hash, "lexemes_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()


@assets_router.post("/process_refined/{file_hash}")
async def process_refined(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """
    Refine a document using Marker API for content extraction and formatting.
    """
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="refined_processing",
            metadata={"processor_type": "refined", "run_id": run_id},
        )

        logger.info(f"Starting document refinement for {file_hash}")
        span.event(name="refinement_started")

        update_asset_status(file_hash, "processing_refined")

        api_key = os.getenv("MARKER_API_KEY")
        base_url = "https://www.datalab.to/api/v1/marker"
        headers = {"X-Api-Key": api_key}

        file_path = os.path.join(
            "/app/filestore/raw",
            f"{file_hash}{os.path.splitext(asset['original_name'])[1]}",
        )

        marker_generation = span.generation(
            name="marker_api_processing",
            input=file_path,
            metadata={"file_hash": file_hash, "file_name": asset["original_name"]},
        )
        try:
            with open(file_path, "rb") as f:
                files = {
                    "file": (asset["original_name"], f, asset["file_type"]),
                    "langs": (None, "English"),
                    "force_ocr": (None, False),
                    "paginate": (None, False),
                }
                response = requests.post(base_url, files=files, headers=headers)

            if not response.ok:
                error_msg = f"Marker API request failed: {response.text}"
                handle_error(span, file_hash, error_msg, logger, update_asset_status)
                raise HTTPException(status_code=response.status_code, detail=error_msg)

            initial_data = response.json()
            marker_generation.event(name="marker_api_initial_response")

            max_polls = 300
            poll_interval = 2
            request_check_url = initial_data["request_check_url"]

            for attempt in range(max_polls):
                time.sleep(poll_interval)
                response = requests.get(request_check_url, headers=headers)
                data = response.json()

                if data["status"] == "complete":
                    if not data["success"]:
                        error_msg = f"Marker processing failed: {data.get('error', 'Unknown error')}"
                        handle_error(
                            marker_generation,
                            file_hash,
                            error_msg,
                            logger,
                            update_asset_status,
                        )
                        raise HTTPException(status_code=500, detail=error_msg)

                    marker_generation.event(name="marker_api_processing_complete")
                    break

                if attempt % 10 == 0:
                    logger.debug(
                        f"Still waiting for results. Attempt {attempt + 1}/{max_polls}"
                    )
                    marker_generation.event(
                        name="marker_api_polling", metadata={"attempt": attempt + 1}
                    )

            if data["status"] != "complete":
                error_msg = f"Marker processing timed out after {max_polls} attempts"
                handle_error(
                    marker_generation, file_hash, error_msg, logger, update_asset_status
                )
                raise HTTPException(status_code=500, detail=error_msg)

        finally:
            marker_generation.end(
                output=data,
            )

        processed_dir = os.path.join("/app/filestore/processed", file_hash)
        os.makedirs(processed_dir, exist_ok=True)

        markdown_path = os.path.join(processed_dir, "content.md")
        with open(markdown_path, "w", encoding="utf-8") as f:
            f.write(data["markdown"])

        meta_path = os.path.join(processed_dir, "meta.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(data["meta"], f, ensure_ascii=False, indent=2)

        update_data = {
            "processed_paths.markdown": markdown_path,
            "processed_paths.meta": meta_path,
            "page_count": data.get("page_count", 1),
            "status": "refined_complete",
        }
        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        logger.info(f"Successfully refined document {file_hash}")
        span.event(
            name="refinement_complete",
            metadata={
                "markdown_path": markdown_path,
                "meta_path": meta_path,
                "page_count": data.get("page_count", 1),
            },
        )

        return {
            "status": "success",
            "markdown_path": markdown_path,
            "meta_path": meta_path,
            "page_count": data.get("page_count", 1),
        }

    except Exception as e:
        error_msg = f"Error refining document: {str(e)}"
        if "span" in locals():
            handle_error(span, file_hash, error_msg, logger, update_asset_status)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()
            time.sleep(0.5)


@assets_router.post("/process_refined_metadata/{file_hash}")
async def process_refined_metadata(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Extract metadata from document content"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="metadata_processing",
            metadata={"processor_type": "metadata", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_metadata")
        logger.info(f"Starting metadata extraction for {file_hash}")

        # Get the processed markdown content
        processed_paths = asset.get("processed_paths", {})
        if not processed_paths or "markdown" not in processed_paths:
            raise HTTPException(
                status_code=400, detail="Markdown content not available"
            )

        # Read the prompt template
        prompts_dir = os.path.join("/app", "prompts", "assets")
        with open(os.path.join(prompts_dir, "metadata.txt"), "r") as f:
            prompt_template = f.read()

        # Read the processed content
        with open(processed_paths["markdown"], "r") as f:
            file_content = f.read()

        # Prepare and make the chat API call
        prompt = prompt_template + "\n\nDocument Content:\n" + file_content
        response = chat_call(query=prompt, messages=[])

        if "error" in response:
            raise HTTPException(status_code=500, detail=response["error"])

        response_text = response["message"]
        json_start = response_text.find("{")
        json_end = response_text.rfind("}") + 1

        if json_start < 0 or json_end <= json_start:
            raise HTTPException(
                status_code=500, detail="Could not find JSON in chat response"
            )

        metadata = json.loads(response_text[json_start:json_end])

        metadata_path = os.path.join(
            "/app/filestore/processed", file_hash, "metadata.json"
        )
        os.makedirs(os.path.dirname(metadata_path), exist_ok=True)

        with open(metadata_path, "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

        # Update asset
        update_data = {
            "metadata": metadata,
            "processed_paths.metadata": metadata_path,
        }
        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        update_asset_status(file_hash, "metadata_complete")
        span.event(name="metadata_complete", metadata={"metadata_path": metadata_path})

        return {
            "status": "success",
            "metadata": metadata,
            "metadata_path": metadata_path,
        }

    except Exception as e:
        error_msg = f"Error extracting metadata: {str(e)}"
        if "span" in locals():
            span.event(
                name="metadata_error", metadata={"error": error_msg}, level="error"
            )
        update_asset_status(file_hash, "metadata_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()


@assets_router.post("/process_refined_splitting/{file_hash}")
async def process_refined_splitting(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Process document content into logical segments"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="splitting_processing",
            metadata={"processor_type": "splitting", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_splitting")
        logger.info(f"Starting content splitting for {file_hash}")

        processed_paths = asset.get("processed_paths", {})
        if not processed_paths or "markdown" not in processed_paths:
            error_msg = "Markdown content not available - need refined content first"
            logger.error(error_msg)
            raise HTTPException(status_code=400, detail=error_msg)

        with open(processed_paths["markdown"], "r") as f:
            file_content = f.read()
            content_length = len(file_content)
            logger.info(f"Document length: {content_length} characters")

        MIN_CHARS_FOR_SPLIT = 24000  # ~6k tokens
        IDEAL_SEGMENT_SIZE = 16000  # ~4k tokens
        MAX_SEGMENT_SIZE = 28000  # ~7k tokens

        if content_length <= MIN_CHARS_FOR_SPLIT:
            logger.info("Document below split threshold - no splitting needed")
            results = {
                "splitRecommendations": {
                    "shouldSplit": False,
                    "confidence": 95,
                    "reasoning": "Document small enough for single-unit processing",
                },
                "documentStats": {
                    "totalLength": {"value": content_length, "unit": "characters"}
                },
            }
            prompt = None
        else:
            prompts_dir = os.path.join("/app", "prompts", "assets")
            with open(os.path.join(prompts_dir, "splitting.txt"), "r") as f:
                prompt_template = f.read()

            prompt = prompt_template + "\n\nDocument Content:\n" + file_content

            generation = span.generation(
                name="splitting_analysis",
                input=prompt,
            )

            response = chat_call(query=prompt, messages=[])

            if "error" in response:
                raise HTTPException(status_code=500, detail=response["error"])

            response_text = response["message"]
            generation.end(output=response_text)

            # Extract JSON from response
            json_start = response_text.find("{")
            json_end = response_text.rfind("}") + 1

            if json_start < 0 or json_end <= json_start:
                raise HTTPException(
                    status_code=500, detail="Could not find JSON in chat response"
                )

            results = json.loads(response_text[json_start:json_end])

        # Validate and adjust split recommendations based on size
        should_split = content_length > MIN_CHARS_FOR_SPLIT
        if should_split and "splitRecommendations" in results:
            splits = results["splitRecommendations"].get("recommendedSplits", [])

            # Validate split sizes
            oversized_splits = []
            total_chars = 0

            for split in splits:
                length = split.get("estimatedLength", {}).get("value", 0)
                if length > MAX_SEGMENT_SIZE:
                    oversized_splits.append(split["suggestedTitle"])
                total_chars += length

            # Add warnings if needed
            if oversized_splits:
                results["splitRecommendations"]["warnings"] = {
                    "oversizedSegments": oversized_splits,
                    "recommendedMaxSize": MAX_SEGMENT_SIZE,
                    "recommendation": "Consider further splitting large segments",
                }

            # Validate total size
            if abs(total_chars - content_length) > content_length * 0.1:
                logger.warning(
                    f"Split size mismatch. Original: {content_length}, "
                    f"Sum of splits: {total_chars}"
                )
                results["splitRecommendations"]["warnings"] = {
                    **results["splitRecommendations"].get("warnings", {}),
                    "sizeMismatch": {
                        "originalSize": content_length,
                        "splitTotalSize": total_chars,
                        "recommendation": "Review split boundaries",
                    },
                }

        # Save results
        processed_dir = os.path.join("/app/filestore/processed", file_hash)
        splitting_path = os.path.join(processed_dir, "splitting.json")
        os.makedirs(os.path.dirname(splitting_path), exist_ok=True)

        with open(splitting_path, "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)

        # Update asset record
        update_data = {
            "splitting": results,
            "processed_paths.splitting": splitting_path,
            "segment_count": len(
                results.get("splitRecommendations", {}).get("recommendedSplits", [])
            ),
            "should_split": results.get("splitRecommendations", {}).get(
                "shouldSplit", False
            ),
        }

        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        update_asset_status(file_hash, "splitting_complete")

        span.event(name="splitting_complete", metadata={"results": results})

        return {
            "status": "success",
            "results": results,
            "splitting_path": splitting_path,
        }

    except Exception as e:
        error_msg = f"Error in content splitting: {str(e)}"
        logger.error(error_msg)
        if "span" in locals():
            span.event(
                name="splitting_error", metadata={"error": error_msg}, level="error"
            )
        update_asset_status(file_hash, "splitting_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()
            time.sleep(0.5)


@assets_router.post("/process_tables/{file_hash}")
async def process_tables(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Process tables from a DOCX file"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        if (
            asset["file_type"]
            != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            logger.info(
                f"Asset {file_hash} is not a DOCX file, skipping table processing"
            )
            update_asset_status(file_hash, "tables_skipped")
            return {"status": "skipped", "reason": "not_docx"}

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="table_processing",
            metadata={"processor_type": "tables", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_tables")
        logger.info(f"Starting table processing for {file_hash}")

        processed_dir = os.path.join("/app/filestore/processed", file_hash)
        tables_dir = os.path.join(processed_dir, "tables")
        os.makedirs(tables_dir, exist_ok=True)

        doc = Document(os.path.join("/app/filestore/raw", f"{file_hash}.docx"))
        table_paths = {}
        tables_meta = {}

        for i, table in enumerate(doc.tables):
            try:
                table_name = f"table_{i}"
                table_data = []

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        text = " ".join(text.split())
                        row_data.append(text)
                    if any(cell for cell in row_data):
                        table_data.append(row_data)

                if not table_data:
                    continue

                # Save as CSV
                csv_path = os.path.join(tables_dir, f"{table_name}.csv")
                with open(csv_path, "w", newline="", encoding="utf-8") as f:
                    import csv

                    writer = csv.writer(f)
                    writer.writerows(table_data)

                # Save as HTML
                html_path = os.path.join(tables_dir, f"{table_name}.html")
                with open(html_path, "w", encoding="utf-8") as f:
                    html = ['<table border="1" class="table">']
                    if table_data:
                        html.append("<thead><tr>")
                        for header in table_data[0]:
                            html.append(f"<th>{header}</th>")
                        html.append("</tr></thead>")
                    html.append("<tbody>")
                    for row in table_data[1:]:
                        html.append("<tr>")
                        for cell in row:
                            html.append(f"<td>{cell}</td>")
                        html.append("</tr>")
                    html.append("</tbody></table>")
                    f.write("\n".join(html))

                table_paths[table_name] = {"csv": csv_path, "html": html_path}

                # Collect metadata
                tables_meta[table_name] = {
                    "num_rows": len(table_data),
                    "num_cols": len(table_data[0]) if table_data else 0,
                    "headers": table_data[0] if table_data else [],
                    "empty_cells": sum(
                        1 for row in table_data for cell in row if not cell.strip()
                    ),
                    "total_cells": sum(len(row) for row in table_data),
                }

            except Exception as e:
                logger.error(f"Error processing table {i}: {str(e)}")
                continue

        # Save metadata
        meta_path = os.path.join(tables_dir, "tables_meta.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(tables_meta, f, ensure_ascii=False, indent=2)

        # Update asset record
        update_data = {
            "has_tables": len(table_paths) > 0,
            "table_count": len(table_paths),
            "processed_paths.tables": table_paths,
        }
        db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        logger.info(f"Completed table processing for {file_hash}")
        update_asset_status(file_hash, "tables_complete")

        span.event(
            name="table_processing_complete",
            metadata={"table_count": len(table_paths)},
        )

        return {
            "status": "success",
            "table_count": len(table_paths),
            "table_paths": table_paths,
            "tables_meta": tables_meta,
        }

    except Exception as e:
        error_msg = f"Error processing tables: {str(e)}"
        if "span" in locals():
            span.event(
                name="table_processing_error",
                metadata={"error": error_msg},
                level="error",
            )
        update_asset_status(file_hash, "tables_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()


@assets_router.post("/process_table_metadata/{file_hash}")
async def process_table_metadata(
    file_hash: str,
    x_span_id: Optional[str] = Header(None),
    x_run_id: Optional[str] = Header(None),
) -> Dict:
    """Process metadata for extracted tables"""
    try:
        db = init_mongo()
        if isinstance(db, dict) and "error" in db:
            raise HTTPException(status_code=500, detail=db["error"])

        asset = db["raw_assets"].find_one({"file_hash": file_hash})
        if not asset:
            raise HTTPException(status_code=404, detail="Asset not found")

        processed_paths = asset.get("processed_paths", {})
        tables = processed_paths.get("tables", {})

        if not tables:
            logger.info(f"No tables found for {file_hash}, skipping metadata")
            update_asset_status(file_hash, "table_metadata_skipped")
            return {"status": "skipped", "reason": "no_tables"}

        run_id = x_run_id or asset.get("current_run_id")
        if not run_id:
            raise HTTPException(status_code=400, detail="No run_id found for asset")

        trace = langfuse.trace(
            name="asset-processing",
            id=run_id,
            metadata={
                "file_hash": file_hash,
                "file_name": asset["original_name"],
                "file_type": asset["file_type"],
                "file_size": asset["file_size"],
            },
        )

        span = trace.span(
            id=x_span_id,
            name="table_metadata_processing",
            metadata={"processor_type": "table_metadata", "run_id": run_id},
        )

        update_asset_status(file_hash, "processing_table_metadata")
        logger.info(f"Starting table metadata processing for {file_hash}")

        # Read the prompt template
        prompts_dir = os.path.join("/app", "prompts", "assets", "table")
        with open(os.path.join(prompts_dir, "metadata.txt"), "r") as f:
            prompt_template = f.read()

        table_metadata = {}

        for table_name, paths in tables.items():
            csv_path = paths.get("csv")
            if not csv_path or not os.path.exists(csv_path):
                continue

            with open(csv_path, "r", encoding="utf-8") as f:
                table_content = f.read()

            prompt = f"{prompt_template}\n\nTable Content:\n{table_content}"
            response = chat_call(query=prompt, messages=[])

            if "error" in response:
                raise HTTPException(
                    status_code=500,
                    detail=f"Chat API error for table {table_name}",
                )

            response_text = response["message"]
            json_start = response_text.find("{")
            json_end = response_text.rfind("}") + 1

            if json_start < 0 or json_end <= json_start:
                logger.error(
                    f"Could not find JSON in chat response for table {table_name}"
                )
                continue

            metadata = json.loads(response_text[json_start:json_end])
            table_metadata[table_name] = metadata

        if table_metadata:
            metadata_path = os.path.join(
                "/app/filestore/processed", file_hash, "tables", "table_metadata.json"
            )
            os.makedirs(os.path.dirname(metadata_path), exist_ok=True)

            with open(metadata_path, "w", encoding="utf-8") as f:
                json.dump(table_metadata, f, ensure_ascii=False, indent=2)

            update_data = {
                "processed_paths.table_metadata": metadata_path,
                "table_metadata": table_metadata,
            }
            db["raw_assets"].update_one({"file_hash": file_hash}, {"$set": update_data})

        update_asset_status(file_hash, "table_metadata_complete")
        span.event(
            name="table_metadata_complete",
            metadata={"table_count": len(table_metadata)},
        )

        return {
            "status": "success",
            "table_metadata": table_metadata,
            "table_count": len(table_metadata),
        }

    except Exception as e:
        error_msg = f"Error processing table metadata: {str(e)}"
        if "span" in locals():
            span.event(
                name="table_metadata_error",
                metadata={"error": error_msg},
                level="error",
            )
        update_asset_status(file_hash, "table_metadata_error", error=error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

    finally:
        if "span" in locals():
            span.end()
